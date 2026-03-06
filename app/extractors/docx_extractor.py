"""
DOCX (Microsoft Word) text extraction using python-docx.
"""

import logging
from pathlib import Path
from typing import Dict, Any
from datetime import datetime

logger = logging.getLogger(__name__)


def extract_docx(file_path: str, document_id: str) -> Dict[str, Any]:
    """
    Extract text from DOCX file using python-docx.
    
    Args:
        file_path: Path to DOCX file
        document_id: Document UUID (as string)
        
    Returns:
        Extraction result dictionary with same structure as PDF extractor
    """
    from docx import Document
    
    logger.info(f"Extracting DOCX: {file_path}")
    
    try:
        doc = Document(file_path)
        paragraphs = []
        full_text_parts = []
        
        # Extract all paragraphs
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:  # Only include non-empty paragraphs
                paragraphs.append({
                    "paragraph_no": i + 1,
                    "text": text,
                    "char_count": len(text)
                })
                full_text_parts.append(text)
        
        # Extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text:
                    table_texts.append(row_text)
        
        if table_texts:
            full_text_parts.append("\n--- Tables ---\n")
            full_text_parts.extend(table_texts)
        
        # Combine all text
        full_text = "\n".join(full_text_parts)
        
        # Estimate "pages" (250 words ≈ 1 page)
        words = full_text.split()
        estimated_pages = max(1, len(words) // 250)
        
        result = {
            "document_id": document_id,
            "filename": Path(file_path).name,
            "extracted_at": datetime.utcnow().isoformat(),
            "text": full_text,
            "pages": paragraphs[:100],  # Limit to first 100 paragraphs in metadata
            "metadata": {
                "num_pages": estimated_pages,
                "is_scanned": False,
                "language": "en",
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "extraction_method": "python-docx",
                "char_count": len(full_text),
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables)
            }
        }
        
        logger.info(f"DOCX extraction complete: {len(paragraphs)} paragraphs, {len(full_text)} chars")
        return result
        
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}", exc_info=True)
        raise Exception(f"Failed to extract DOCX: {str(e)}")
